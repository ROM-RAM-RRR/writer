from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from openai import OpenAI
from typing import Optional
import os
from app.services.config_service import config_service
from app.services.theme_service import theme_service
from app.services.outline_service import outline_service
from app.models.outline import OutlineCreate

router = APIRouter()

# 禁用代理
os.environ.pop('http_proxy', None)
os.environ.pop('https_proxy', None)
os.environ.pop('HTTP_PROXY', None)
os.environ.pop('HTTPS_PROXY', None)


class OutlineRequest(BaseModel):
    theme: str
    genre: Optional[str] = None
    style: Optional[str] = None
    length: Optional[str] = "中篇小说"
    requirements: Optional[str] = None
    theme_id: Optional[str] = None
    num_options: Optional[int] = 3


@router.post("/")
async def generate_outline(request: OutlineRequest):
    config = config_service.get_config()

    if not config.get("api_key"):
        raise HTTPException(status_code=400, detail="API key not configured")

    # 构建系统提示词
    system_prompt = """你是一个小说大纲创作专家。根据用户提供的素材，创作3-5个不同风格的故事大纲。

每个大纲应包含：
1. 故事标题
2. 故事类型/风格
3. 核心情节线
4. 主要人物设定（2-4个）
5. 故事亮点/卖点
6. 预计章节数或篇幅

请以JSON数组格式返回，格式如下：
[
  {
    "title": "故事标题",
    "genre": "故事类型",
    "plot": "核心情节线",
    "characters": ["人物1", "人物2", ...],
    "highlights": "故事亮点",
    "chapters": 章节数
  },
  ...
]

确保每个大纲都有独特的创意和差异化。"""

    user_prompt = f"""请为以下主题创作小说大纲：

主题：{request.theme}
类型：{request.genre or "未指定"}
风格：{request.style or "未指定"}
篇幅：{request.length or "中篇小说"}
"""

    if request.requirements:
        user_prompt += f"\n其他要求：{request.requirements}"

    # 如果有主题ID，添加主题设定
    if request.theme_id:
        theme = theme_service.get_theme(request.theme_id)
        if theme:
            user_prompt += f"\n\n主题设定：\n类型：{theme.genre}\n风格：{theme.style}\n背景：{theme.background}\n角色：{theme.characters}"

    client = OpenAI(
        base_url=config.get("base_url", "https://api.openai.com/v1"),
        api_key=config["api_key"],
    )

    model = config.get("model", "gpt-3.5-turbo")

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.8,
            max_tokens=4000,
        )

        content = response.choices[0].message.content

        # 尝试解析JSON
        import json
        import re

        # 提取JSON数组
        json_match = re.search(r'\[.*\]', content, re.DOTALL)
        if json_match:
            outlines = json.loads(json_match.group())
            return {"outlines": outlines, "raw": content}
        else:
            return {"outlines": [], "raw": content}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"生成大纲失败: {str(e)}")


# 继续生成/修改大纲
class ContinueRequest(BaseModel):
    selected_outline: dict
    action: str  # "continue" 或 "modify"
    modifications: Optional[str] = None
    theme_id: Optional[str] = None
    max_tokens: Optional[int] = 2000
    temperature: Optional[float] = 0.8


@router.post("/continue")
async def continue_outline(request: ContinueRequest):
    config = config_service.get_config()

    if not config.get("api_key"):
        raise HTTPException(status_code=400, detail="API key not configured")

    system_prompt = """你是一个小说写作助手。根据用户选择的剧情大纲，继续发展故事内容。

如果用户选择"继续"，请在选定的大纲基础上，生成故事的开篇章节内容（约1000-2000字）。
如果用户选择"修改"，请根据用户的修改要求，重新调整大纲并给出新的版本。

返回格式：
{
  "type": "continue" 或 "modify",
  "content": "生成的内容或修改后的大纲"
}"""

    user_content = f"选择的剧情大纲：\n{request.selected_outline}\n\n"

    if request.action == "continue":
        user_content += "请继续生成故事的开篇章节内容。"
    elif request.action == "modify":
        user_content += f"请根据以下要求修改大纲：{request.modifications}"

    # 添加主题设定
    if request.theme_id:
        theme = theme_service.get_theme(request.theme_id)
        if theme:
            user_content += f"\n\n主题设定：\n类型：{theme.genre}\n风格：{theme.style}\n背景：{theme.background}\n角色：{theme.characters}"

    client = OpenAI(
        base_url=config.get("base_url", "https://api.openai.com/v1"),
        api_key=config["api_key"],
    )

    model = config.get("model", "gpt-3.5-turbo")

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=request.temperature or 0.8,
            max_tokens=request.max_tokens or 2000,
        )

        return {"content": response.choices[0].message.content}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"操作失败: {str(e)}")


# 保存大纲
class SaveOutlineRequest(BaseModel):
    title: str
    genre: str
    plot: str
    characters: list[str]
    highlights: str
    chapters: int
    source_theme: Optional[str] = None
    content: Optional[str] = None
    project_id: Optional[str] = None


@router.post("/save")
async def save_outline(request: SaveOutlineRequest):
    try:
        outline = outline_service.create_outline(
            OutlineCreate(
                title=request.title,
                genre=request.genre,
                plot=request.plot,
                characters=request.characters,
                highlights=request.highlights,
                chapters=request.chapters,
                source_theme=request.source_theme,
                content=request.content,
            ),
            project_id=request.project_id,
        )
        return outline
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"保存大纲失败: {str(e)}")


# 获取大纲列表
@router.get("/")
async def list_outlines(project_id: Optional[str] = None):
    outlines = outline_service.list_outlines(project_id)
    return outlines


# 获取单个大纲
@router.get("/{outline_id}")
async def get_outline(outline_id: str):
    outline = outline_service.get_outline(outline_id)
    if not outline:
        raise HTTPException(status_code=404, detail="大纲不存在")
    return outline


# 删除大纲
@router.delete("/{outline_id}")
async def delete_outline(outline_id: str):
    if not outline_service.delete_outline(outline_id):
        raise HTTPException(status_code=404, detail="大纲不存在")
    return {"message": "大纲已删除"}


# 导出大纲为Word
@router.get("/{outline_id}/export")
async def export_outline(outline_id: str):
    from fastapi.responses import FileResponse
    from docx import Document
    from docx.shared import Inches, Pt
    import io

    outline = outline_service.get_outline(outline_id)
    if not outline:
        raise HTTPException(status_code=404, detail="大纲不存在")

    # 创建Word文档
    doc = Document()

    # 标题
    title = doc.add_heading(outline.title, level=1)
    title.alignment = 1  # 居中

    # 基本信息
    doc.add_heading('基本信息', level=2)
    doc.add_paragraph(f'类型：{outline.genre}')
    doc.add_paragraph(f'预计章节数：{outline.chapters}')
    doc.add_paragraph(f'创建时间：{outline.created_at}')

    if outline.source_theme:
        doc.add_paragraph(f'来源主题：{outline.source_theme}')

    # 情节
    doc.add_heading('情节概要', level=2)
    doc.add_paragraph(outline.plot)

    # 角色
    doc.add_heading('主要角色', level=2)
    for char in outline.characters:
        doc.add_paragraph(f'• {char}')

    # 亮点
    doc.add_heading('故事亮点', level=2)
    doc.add_paragraph(outline.highlights)

    # 如果有生成的正文内容
    if outline.content:
        doc.add_heading('已生成内容', level=2)
        doc.add_paragraph(outline.content)

    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return FileResponse(
        buffer,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=f'{outline.title}.docx'
    )