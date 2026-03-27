from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from app.models.project import Project, ProjectCreate, ProjectUpdate
from app.services.project_service import project_service
from docx import Document
import io

router = APIRouter()


@router.get("/", response_model=list[Project])
async def list_projects():
    return project_service.list_projects()


@router.get("/{project_id}", response_model=Project)
async def get_project(project_id: str):
    project = project_service.get_project(project_id)
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@router.post("/", response_model=Project)
async def create_project(project: ProjectCreate):
    return project_service.create_project(project)


@router.put("/{project_id}", response_model=Project)
async def update_project(project_id: str, project: ProjectUpdate):
    updated = project_service.update_project(project_id, project)
    if not updated:
        raise HTTPException(status_code=404, detail="Project not found")
    return updated


@router.delete("/{project_id}")
async def delete_project(project_id: str):
    if not project_service.delete_project(project_id):
        raise HTTPException(status_code=404, detail="Project not found")
    return {"message": "Project deleted"}


# 导出项目为Word
@router.get("/{project_id}/export")
async def export_project(project_id: str):
    try:
        project = project_service.get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="Project not found")

        # 创建Word文档
        doc = Document()

        # 标题
        title = doc.add_heading(project.title, level=1)
        title.alignment = 1

        # 元信息
        doc.add_paragraph(f'创建时间：{project.created_at}')
        doc.add_paragraph(f'更新时间：{project.updated_at}')
        doc.add_paragraph(f'字数：约 {len(project.content)} 字')

        # 大纲（如果有）
        if project.outline:
            doc.add_heading('故事大纲', level=2)
            doc.add_paragraph(project.outline)

        # 正文内容
        doc.add_heading('正文', level=2)

        # 按段落分隔内容
        paragraphs = project.content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)

        # 保存到文件
        import tempfile
        import os
        from urllib.parse import quote

        # 使用ASCII安全的文件名
        safe_filename = f"project_{project_id}.docx"
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False, mode='wb') as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name

        # 读取并删除临时文件
        with open(tmp_path, 'rb') as f:
            content = f.read()
        os.unlink(tmp_path)

        from starlette.responses import Response
        return Response(
            content,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={'Content-Disposition': f'attachment; filename="{safe_filename}"'}
        )
    except Exception as e:
        import traceback
        raise HTTPException(status_code=500, detail=f"导出失败: {str(e)}")